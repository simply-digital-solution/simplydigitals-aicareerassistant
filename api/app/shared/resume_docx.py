"""Build a .docx from a GeneratedResumeOutput using python-docx.

Mirrors the styling in TailoredResumePanel.tsx / downloadAsDocx:
- Calibri 11pt body
- 28pt bold name, left-aligned
- Header lines in gray below name
- Section headings: blue uppercase, bottom border
- Experience: bold title + italic company + tab-right dates, then bullets
"""
from io import BytesIO

from docx import Document
from docx.shared import Pt, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from app.shared.schemas import GeneratedResumeOutput

HEADING_BLUE = RGBColor(0x1F, 0x5C, 0x9E)
GRAY_TEXT = RGBColor(0x55, 0x55, 0x55)
DATE_GRAY = RGBColor(0x88, 0x88, 0x88)
BODY_FONT = "Calibri"
BODY_PT = 11
NAME_PT = 28


def _set_run(run, text: str, *, bold=False, italic=False, size_pt=BODY_PT,
             color: RGBColor | None = None, font: str = BODY_FONT):
    run.text = text
    run.font.name = font
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def _add_bottom_border(paragraph, color_hex: str = "1F5C9E"):
    """Add a bottom border to a paragraph via OOXML."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _page_width_twips() -> int:
    # A4 minus 0.75" left + 0.75" right margins: (11906 - 2*1080) = 9746 twips
    # Use same value as frontend: 9360 (≈ 1" margins)
    return 9360


def build_docx_bytes(resume: GeneratedResumeOutput) -> bytes:
    doc = Document()

    # Page margins: 0.5" top/bottom, 0.75" sides (matches frontend)
    for section in doc.sections:
        section.top_margin = Twips(720)
        section.bottom_margin = Twips(720)
        section.left_margin = Twips(1080)
        section.right_margin = Twips(1080)

    # Default paragraph spacing: remove extra space after
    doc.styles["Normal"].paragraph_format.space_after = Pt(0)

    # ── Candidate name ──────────────────────────────────────────────────────
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_name.paragraph_format.space_after = Pt(2)
    r = p_name.add_run(resume.name)
    _set_run(r, resume.name, bold=True, size_pt=NAME_PT)

    # ── Header lines (verbatim) ─────────────────────────────────────────────
    for i, line in enumerate(resume.header_lines or []):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        is_last = i == len(resume.header_lines) - 1
        p.paragraph_format.space_after = Pt(10 if is_last else 2)
        r = p.add_run(line)
        _set_run(r, line, color=GRAY_TEXT)

    # ── Sections ────────────────────────────────────────────────────────────
    for section in resume.sections:
        # Section heading
        p_heading = doc.add_paragraph()
        p_heading.paragraph_format.space_before = Pt(14)
        p_heading.paragraph_format.space_after = Pt(5)
        _add_bottom_border(p_heading)
        r = p_heading.add_run(section.title.upper())
        _set_run(r, section.title.upper(), bold=False, color=HEADING_BLUE)

        if section.section_type == "experience":
            for entry in section.experience:
                # Title + company + dates line
                p_title = doc.add_paragraph()
                p_title.paragraph_format.space_before = Pt(6)
                p_title.paragraph_format.space_after = Pt(2)
                # Right-aligned dates via tab stop
                tab_stop = OxmlElement("w:tab")
                pPr = p_title._p.get_or_add_pPr()
                tabs_el = OxmlElement("w:tabs")
                tab_el = OxmlElement("w:tab")
                tab_el.set(qn("w:val"), "right")
                tab_el.set(qn("w:pos"), str(_page_width_twips()))
                tabs_el.append(tab_el)
                pPr.append(tabs_el)

                r1 = p_title.add_run(entry.title)
                _set_run(r1, entry.title, bold=True)
                r2 = p_title.add_run("  ")
                _set_run(r2, "  ")
                r3 = p_title.add_run(entry.company)
                _set_run(r3, entry.company, italic=True, color=GRAY_TEXT)
                r_tab = p_title.add_run("\t")
                _set_run(r_tab, "\t")
                r4 = p_title.add_run(entry.dates)
                _set_run(r4, entry.dates, color=DATE_GRAY)

                # 2-line role summary (job-relevant highlights)
                if entry.summary:
                    p_sum = doc.add_paragraph()
                    p_sum.paragraph_format.space_before = Pt(3)
                    p_sum.paragraph_format.space_after = Pt(3)
                    r_sum = p_sum.add_run(entry.summary)
                    _set_run(r_sum, entry.summary, color=GRAY_TEXT)

                for bullet in entry.bullets:
                    p_b = doc.add_paragraph(style="List Bullet")
                    p_b.paragraph_format.space_before = Pt(1)
                    p_b.paragraph_format.space_after = Pt(1)
                    r = p_b.add_run(bullet)
                    _set_run(r, bullet)
        else:
            for line in section.content:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                r = p.add_run(line)
                _set_run(r, line)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
