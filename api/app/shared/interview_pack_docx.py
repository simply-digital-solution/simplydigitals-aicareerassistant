"""Build a .docx for an interview pack (pitch + STAR questions) using python-docx."""
from io import BytesIO

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from app.shared.schemas import InterviewPackOutput

HEADING_BLUE = RGBColor(0x1F, 0x5C, 0x9E)
BODY_FONT = "Calibri"
BODY_PT = 11
TITLE_PT = 22
SECTION_PT = 13


def _set_run(run, text: str, *, bold=False, italic=False, size_pt=BODY_PT,
             color: RGBColor | None = None):
    run.text = text
    run.bold = bold
    run.italic = italic
    run.font.name = BODY_FONT
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = color
    # theme colour override
    rpr = run._r.get_or_add_rPr()
    cs = OxmlElement('w:rFonts')
    cs.set(qn('w:cs'), BODY_FONT)
    rpr.append(cs)


def _add_section_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text.upper())
    _set_run(run, text.upper(), bold=True, size_pt=SECTION_PT, color=HEADING_BLUE)
    # bottom border
    ppr = p._p.get_or_add_pPr()
    pb = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F5C9E')
    pb.append(bottom)
    ppr.append(pb)


def build_interview_pack_docx(
    pack: InterviewPackOutput,
    company: str,
    job_title: str,
) -> bytes:
    doc = Document()

    # Margins
    for section in doc.sections:
        section.top_margin = Pt(48)
        section.bottom_margin = Pt(48)
        section.left_margin = Pt(72)
        section.right_margin = Pt(72)

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = title_p.add_run(f"Interview Prep — {company}")
    _set_run(r, f"Interview Prep — {company}", bold=True, size_pt=TITLE_PT)

    sub_p = doc.add_paragraph()
    r2 = sub_p.add_run(job_title)
    _set_run(r2, job_title, size_pt=BODY_PT, color=RGBColor(0x55, 0x55, 0x55))
    sub_p.paragraph_format.space_after = Pt(8)

    # 2-Minute Pitch
    _add_section_heading(doc, "2-Minute Pitch")
    pitch_p = doc.add_paragraph()
    pitch_p.paragraph_format.space_after = Pt(6)
    r3 = pitch_p.add_run(pack.pitch)
    _set_run(r3, pack.pitch, size_pt=BODY_PT)

    # STAR Questions
    _add_section_heading(doc, f"STAR Questions ({len(pack.star_questions)})")

    for i, q in enumerate(pack.star_questions, 1):
        # Question heading
        q_p = doc.add_paragraph()
        q_p.paragraph_format.space_before = Pt(10)
        q_p.paragraph_format.space_after = Pt(2)
        r_num = q_p.add_run(f"{i}. ")
        _set_run(r_num, f"{i}. ", bold=True, size_pt=BODY_PT, color=HEADING_BLUE)
        r_q = q_p.add_run(q.q)
        _set_run(r_q, q.q, bold=True, size_pt=BODY_PT)

        for label, value in [
            ("Situation", q.situation),
            ("Task", q.task),
            ("Action", q.action),
            ("Result", q.result),
        ]:
            row_p = doc.add_paragraph()
            row_p.paragraph_format.left_indent = Pt(18)
            row_p.paragraph_format.space_after = Pt(1)
            r_label = row_p.add_run(f"{label}: ")
            _set_run(r_label, f"{label}: ", bold=True, size_pt=BODY_PT)
            r_val = row_p.add_run(value)
            _set_run(r_val, value, size_pt=BODY_PT)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
